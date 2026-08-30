from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT

from common.report_utils import create_styled_document, add_title


def generate_report(
    output_path: Path,
    stage1_stats: dict | None = None,
    stage2_stats: dict | None = None,
    total_rows: int = 0,
    elapsed_total: float = 0,
    cost_stage1: str = "N/A",
    cost_stage2: str = "N/A",
):
    doc = create_styled_document()

    add_title(doc, "Кейс №3: Ревизор кода")

    doc.add_heading("1. Описание решения", level=1)
    doc.add_paragraph(
        "Пайплайн статического ревью кода на уязвимости с использованием двух "
        "уровней LLM-анализа через Yandex Cloud AI Studio (OpenAI-совместимый API)."
    )

    doc.add_heading("Архитектура", level=2)
    doc.add_paragraph(
        "Двухуровневая схема обеспечивает баланс между стоимостью и качеством анализа:"
    )

    items = [
        "Этап 1 — триаж (qwen3-235b-a22b-fp8): быстрый вердикт vulnerable/secure/uncertain для всех ~19 000 строк датасета.",
        "Этап 2 — полный анализ (qwen3-235b-a22b-fp8): детальный разбор (CWE ID, механизм эксплуатации, безопасный фикс, обоснование) только для строк, флагнутых на этапе 1 как vulnerable или uncertain.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Ключевые решения", level=2)
    decisions = [
        "Cache-friendly промпт: статичная инструкция вынесена в поле instructions, переменный код — в input, что позволяет провайдеру кешировать повторяющийся префикс.",
        "Структурированный JSON-вывод с retry-логикой: при невалидном JSON ответ повторяется до 2 раз, после чего строка деградирует в uncertain с флагом parse_error.",
        "Валидация CWE ID: формат CWE-\\d+ плюс проверка по локальному справочнику (1219 ID из официального каталога MITRE). Несуществующие CWE понижают вердикт до uncertain.",
        "Конкурентный батчинг: asyncio.Semaphore с ограничением 10 (триаж) и 10 (полный анализ) одновременных запросов, retry с экспоненциальной задержкой на 429/5xx.",
        "Запрет исполнения кода: ни один компонент пайплайна не вызывает exec/eval/subprocess над кодом из датасета или ответами модели. Подтверждено статическим аудит-тестом.",
    ]
    for d in decisions:
        doc.add_paragraph(d, style="List Number")

    doc.add_heading("2. Инструменты", level=1)
    tools_data = [
        ("Модель", "qwen3-235b-a22b-fp8 (Yandex Cloud AI Studio)"),
        ("API", "OpenAI-совместимый SDK, Responses API (client.responses.create)"),
        ("Справочник CWE", "Локальный JSON-снапшот каталога MITRE, 1219 ID"),
        ("Язык пайплайна", "Python 3.12+, asyncio, pandas, tenacity"),
        ("Тестирование", "pytest, unit + integration + e2e, статический аудит безопасности"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Компонент"
    hdr[1].text = "Решение"
    for name, val in tools_data:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = val

    doc.add_heading("3. Алгоритм работы", level=1)
    steps = [
        "Загрузка датасета data/case_3.csv (~19 000 строк, поля unique_id, code).",
        "Этап 1: триаж каждой строки моделью qwen3-235b-a22b-fp8 с температурой 0.0 (детерминированный вердикт).",
        "Фильтрация: отбор строк с вердиктом vulnerable/uncertain для этапа 2.",
        "Этап 2: полный анализ отобранных строк моделью qwen3-235b-a22b-fp8 — CWE ID, механизм, фикс, обоснование.",
        "Валидация CWE ID по справочнику, понижение вердикта при невалидном CWE.",
        "Объединение результатов, экспорт в .xlsx/.csv с валидацией полноты.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_heading("4. Результаты прогона", level=1)
    if stage1_stats:
        doc.add_paragraph(f"Всего строк: {total_rows}")
        doc.add_paragraph(f"Распределение вердиктов (этап 1):")
        for k, v in sorted(stage1_stats.items()):
            doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        doc.add_paragraph(f"Стоимость этапа 1: {cost_stage1}")

    if stage2_stats:
        doc.add_paragraph(f"Строк на этапе 2: {sum(stage2_stats.values())}")
        for k, v in sorted(stage2_stats.items()):
            doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        doc.add_paragraph(f"Стоимость этапа 2: {cost_stage2}")

    if elapsed_total > 0:
        doc.add_paragraph(f"Общее время выполнения: {elapsed_total:.1f} сек")

    doc.add_heading("5. Ограничения и зоны развития", level=1)

    doc.add_heading("Ограничения", level=2)
    limitations = [
        "Код и предложенные исправления не компилировались и не выполнялись — верификация корректности фикса возможна только через статический анализ и экспертную оценку.",
        "Мультиязычность без явной языковой маршрутизации: модель определяет язык по содержимому сниппета. Качество для редких языков может быть ниже.",
        "Качество CWE-присвоения зависит от модели без специализированного fine-tuning — возможны галлюцинации редких CWE-номеров (снижено валидацией по справочнику).",
        "Ложноотрицательные результаты на этапе 1: если триаж пропустит уязвимость как secure, она не попадёт на этап 2. Компенсируется «осторожным» промптом (при сомнениях — uncertain).",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    doc.add_heading("Неполнота обработки и инфраструктурные ограничения", level=2)
    doc.add_paragraph(
        "Представленные результаты демонстрируют работоспособность пайплайна на обработанной "
        "части выборки из 18 864 сниппетов. В ходе эксплуатации были выявлены ограничения "
        "инфраструктуры Yandex Cloud AI Studio, влияющие на полноту обработки:"
    )
    infra_issues = [
        "При конкурентных запросах (batch-инференс) часть ответов возвращается с пустым "
        "содержимым (HTTP 200, пустой output_text) — модель генерирует внутренний reasoning, "
        "но не формирует текстовый ответ. Это системное ограничение API, а не ошибка пайплайна.",
        "Разные модели имеют различные лимиты конкурентности: от 5 (yandexgpt-5-pro) до 150 "
        "(qwen3-235b) одновременных запросов. Превышение лимита приводит к деградации "
        "качества ответов без явной ошибки HTTP.",
        "Стоимость обработки полного датасета через премиум-модели (yandexgpt-5-pro, qwen3-235b) "
        "существенно выше, чем через более лёгкие open-source модели. "
        "Для production-масштабирования необходим пересмотр баланса между качеством и стоимостью.",
    ]
    for issue in infra_issues:
        doc.add_paragraph(issue, style="List Bullet")

    doc.add_paragraph(
        "Для обеспечения полной обработки датасета с гарантированным качеством необходимо:"
    )
    needs = [
        "Выделение бюджета на API-вызовы: ориентировочно 3 000–5 000 ₽ на полный прогон "
        "двух этапов с retry-логикой для обработки пустых ответов.",
        "Расширение квот Yandex Cloud AI Studio: увеличение лимитов конкурентности для "
        "выбранных моделей, что позволит сократить время обработки с часов до десятков минут.",
        "Резервирование вычислительных ресурсов: dedicated endpoint или SLA-гарантии от "
        "провайдера для стабильности batch-инференса в production-среде.",
    ]
    for need in needs:
        doc.add_paragraph(need, style="List Number")

    doc.add_heading("Вектор развития", level=2)
    improvements = [
        "Adversarial fine-tuning триаж-модели на пропущенных уязвимостях для снижения false negative rate.",
        "Языковая маршрутизация: определение языка сниппета и специализированные промпты для C/C++, Python, PHP и др.",
        "Интеграция с SIEM-системами для мониторинга data drift и деградации качества модели в продакшене.",
        "Добавление heuristic-слоя (regex-паттерны на известные уязвимые конструкции) как дополнительный сигнал к LLM-вердикту.",
        "Online learning: накопление экспертных разметок для периодического дообучения модели.",
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style="List Number")

    doc.save(str(output_path))
    return output_path
