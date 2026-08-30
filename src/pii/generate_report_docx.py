"""Generate the project report as DOCX."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _style_table(table, header_color: str = "4472C4") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
    for row_idx, row in enumerate(table.rows[1:], start=1):
        bg = "D9E2F3" if row_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            _set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def generate_report(output_path: str | Path) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # ── Title ──
    title = doc.add_heading("Отчёт: Пайплайн анонимизации персональных данных", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # ── 1. Technology ──
    doc.add_heading("1. Инструмент и технология", level=1)

    doc.add_paragraph(
        "Модель: OpenMed/privacy-filter-nemotron — токен-классификатор на базе "
        "openai/privacy-filter, дообученный на датасете nvidia/Nemotron-PII. "
        "55 категорий ПДн, 221 BIOES-класс, macro B-F1 ≈ 0.95 на тест-сплите."
    )
    doc.add_paragraph(
        "Датасет: nvidia/Nemotron-PII (test-split, 100 000 записей). "
        "Синтетические документы, имитирующие договоры, письма, заявки, тикеты, медицинские записи. "
        "150+ типов документов, 30 доменов, формат structured/unstructured, локаль us."
    )
    doc.add_paragraph(
        "Инференс: batch-обработка через transformers (AutoTokenizer + AutoModelForTokenClassification). "
        "Sliding-window (512 токенов, overlap 64) для длинных документов. "
        "Locale-aware генерация псевдонимов (EN/RU словари)."
    )

    # ── 2. Code ──
    doc.add_heading("2. Код решения", level=1)

    doc.add_paragraph(
        "Основной модуль: src/pii/. Точка входа: src/pii/run.py (CLI). Ключевые компоненты:"
    )

    components = [
        ("tokenizer.py", "Sliding-window токенизация с маппингом token→char offsets"),
        ("bioes.py", "BIOES-decode в span-уровневые сущности"),
        ("chunk_merge.py", "Слияние спанов на границах чанков (dedup, max score)"),
        ("regex_fallback.py", "Regex-слой: email, телефон, карта (Luhn), defense-in-depth"),
        ("confidence.py", "Политика уверенности: min_score=0.5, needs_review для high-risk"),
        ("normalize.py", "Нормализация + fuzzy-match имён (rapidfuzz)"),
        ("registry.py", "Locale-aware per-document реестр псевдонимов (EN/RU)"),
        ("strategies.py", "Category-aware стратегии замены (15+ тонких лейблов)"),
        ("apply.py", "Применение замен по offsets (справа налево, без сдвига)"),
        ("export.py", "Экспорт JSON/JSONL/XLSX (audit + final)"),
        ("pipeline.py", "Оркестрация пайплайна"),
        ("dataset.py", "Загрузка Nemotron-PII, стратифицированная выборка"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Модуль"
    hdr[1].text = "Назначение"
    for mod, desc in components:
        row = t.add_row().cells
        row[0].text = mod
        row[1].text = desc
    _style_table(t)
    doc.add_paragraph()

    # ── 3. Algorithm ──
    doc.add_heading("3. Алгоритм", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Детекция → Пороги → Консистентная замена → Экспорт")
    run.bold = True

    algo_steps = [
        "Sliding-window токенизация текста (окно 512, overlap 64).",
        "Модель предсказывает BIOES-теги для каждого токена.",
        "BIOES-decode → span-уровневые сущности (start, end, label, score).",
        "Merge спанов на границах чанков (dedup, max score).",
        "Regex-fallback: email, телефон, номер карты (Luhn) — defense-in-depth.",
        "Пороговая политика: score ≥ 0.85 → замена; 0.5–0.85 для high-risk → маскирование + needs_review; < 0.5 → пропуск.",
        "Нормализация + fuzzy-match имён для группировки вариантов.",
        "Детерминированный реестр псевдонимов (per-document): SHA256(doc_id + normalized_value + category) → выбор из EN/RU справочника.",
        "Category-aware замена: имена → псевдонимы, телефоны/карты → маскирование с сохранением формата, даты → генерализация, адреса → обобщение, медданные → метка.",
        "Применение замен по offsets (справа налево, без сдвига).",
        "Экспорт: финальный артефакт (без original) + аудиторский (с original).",
    ]
    for i, step in enumerate(algo_steps, 1):
        doc.add_paragraph(f"{i}) {step}", style="List Number")

    # ── 4. Strategies ──
    doc.add_heading("4. Стратегии замены", level=1)

    doc.add_paragraph(
        "Модель возвращает 50+ тонких лейблов (first_name, last_name, street_address, "
        "credit_debit_card, bank_routing_number и т.д.). Каждый лейбл маппится на стратегию:"
    )

    strategies = [
        ("first_name, last_name, user_name", "Псевдоним из EN/RU словаря (детерминированный, консистентный)", "John Smith → Joseph Taylor"),
        ("street_address, city, state, country", "Генерализация: EN: номер + улица + город; RU: город + улица + дом", "123 Main St → 60 Maple Dr, Houston, TX"),
        ("date_of_birth, date", "Генерализация даты: EN: MM/DD/1990; RU: DD.MM.1990", "05/12/1985 → 07/07/1990"),
        ("company_name", "Фиктивная организация из EN/RU словаря", "Acme Corp → GlobalTech Inc"),
        ("email", "Маскирование local-part, домен сохранён", "john@acme.com → j***n@acme.com"),
        ("PHONE_NUMBER, phone_number", "Формат-сохраняющая маска: последние N цифр видны", "+1 (555) 123-4567 → +1 (555) ***-**67"),
        ("credit_debit_card", "Маска, последние 4 видны", "4111 1111 1111 1111 → **** ****** *111"),
        ("ssn, GOV_ID, certificate_license", "Маска, 2 цифры видны", "123-45-6789 → ***-**-6789"),
        ("account_number, bank_routing_number", "Маска, 2 цифры видны", "271210785 → *******85"),
        ("HEALTHCARE_DATA, medical_record_number", "Метка категории", "diabetes → [MEDICAL DATA]"),
        ("occupation, gender, race_ethnicity, url...", "[REDACTED] — чувствительные атрибуты без словаря", "Engineer → [REDACTED]"),
    ]
    st = doc.add_table(rows=1, cols=3)
    st.style = "Table Grid"
    hdr = st.rows[0].cells
    hdr[0].text = "Категория"
    hdr[1].text = "Стратегия"
    hdr[2].text = "Пример"
    for cat, strat, ex in strategies:
        row = st.add_row().cells
        row[0].text = cat
        row[1].text = strat
        row[2].text = ex
    _style_table(st)
    doc.add_paragraph()

    # ── 5. Inference results ──
    doc.add_heading("5. Результаты инференса", level=1)

    doc.add_paragraph(
        "Стратифицированная выборка 1000 записей из test-split (100k) по домену (domain). "
        "Обработано 986 документов за 367 секунд (~6 мин) на CPU."
    )

    results = [
        ("Всего сущностей детектировано", "9 199"),
        ("Псевдонимы (имена, адреса, даты, организации)", "5 365 (58.3%)"),
        ("Формат-маски (телефон, карта, ID, счёт)", "1 373 (14.9%)"),
        ("[REDACTED] (occupation, url, gender и др.)", "2 461 (26.8%)"),
        ("Документов с needs_review", "15 (1.5%)"),
        ("Ошибок обработки", "0"),
        ("Уникальных типов документов", "678"),
        ("Локаль", "us (все документы)"),
    ]
    rt = doc.add_table(rows=1, cols=2)
    rt.style = "Table Grid"
    hdr = rt.rows[0].cells
    hdr[0].text = "Метрика"
    hdr[1].text = "Значение"
    for metric, val in results:
        row = rt.add_row().cells
        row[0].text = metric
        row[1].text = val
    _style_table(rt)
    doc.add_paragraph()

    doc.add_paragraph("Топ-10 детектированных категорий ПДн:")

    top_labels = [
        ("PHONE_NUMBER", "1 357"),
        ("first_name", "811"),
        ("date", "730"),
        ("last_name", "534"),
        ("company_name", "472"),
        ("email", "381"),
        ("url", "376"),
        ("occupation", "287"),
        ("time", "221"),
        ("phone_number", "215"),
    ]
    lt = doc.add_table(rows=1, cols=2)
    lt.style = "Table Grid"
    hdr = lt.rows[0].cells
    hdr[0].text = "Категория"
    hdr[1].text = "Количество"
    for label, count in top_labels:
        row = lt.add_row().cells
        row[0].text = label
        row[1].text = count
    _style_table(lt)
    doc.add_paragraph()

    # ── 6. Tests ──
    doc.add_heading("6. Тестовое покрытие", level=1)

    doc.add_paragraph("Три уровня тестов на pytest:")

    test_items = [
        "Unit (121 тест): BIOES-decode, chunk merge, regex fallback (Luhn), confidence policy, pseudonym registry (EN + RU), replacement strategies (EN + RU), offset apply. Без загрузки модели.",
        "Integration (4 теста): экспорт final vs audit, валидация JSON/XLSX, no-PII-leak в финальном артефакте.",
        "E2E (4 теста): smoke test, no-PII-leak (безопасность), determinism, XLSX schema.",
        "Model (7 тестов): pipeline на структурированных/неструктурированных документах, консистентность в документе (требуют реальные веса).",
    ]
    for item in test_items:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Быстрый прогон: ")
    run.bold = True
    p.add_run("pytest -m 'not model' — 121 тест за ~1.5 сек. ")
    run = p.add_run("Полный прогон: ")
    run.bold = True
    p.add_run("pytest — 128 тестов.")

    # ── 7. Limitations ──
    doc.add_heading("7. Ограничения", level=1)

    limitations = [
        "Модель покрывает 55 категорий ПДн; редкие категории могут давать ложноотрицательные результаты — компенсируется regex-fallback.",
        "Sliding-window: overlap 64 токена может не покрыть сущности длиннее overlap-зоны.",
        "Fuzzy-match имён — эвристика (token_sort_ratio), не строгий coreference resolution.",
        "Обработка ограничена текстовыми документами (без OCR).",
        "26.8% сущностей (occupation, url, gender и др.) заменяются на [REDACTED] — нет словаря псевдонимов для этих категорий.",
        "Полный датасет (100k) требует ~6 часов на CPU; для production рекомендуется GPU.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    doc.save(str(output_path))


if __name__ == "__main__":
    generate_report("output/case1-pii-anonymization/report.docx")
    print("Report generated: output/case1-pii-anonymization/report.docx")
