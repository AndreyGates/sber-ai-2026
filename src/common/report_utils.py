"""Shared report generation utilities for hackathon case reports."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def create_styled_document() -> Document:
    """Create a Document with standardized styles."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    return doc


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    """Add a centered title with optional subtitle."""
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def style_table(table, header_color: str = "4472C4") -> None:
    """Apply consistent styling to a table: colored header, alternating rows."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
    for row_idx, row in enumerate(table.rows[1:], start=1):
        bg = "D9E2F3" if row_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            _set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              header_color: str = "4472C4") -> None:
    """Add a styled table with headers and data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    style_table(table, header_color)
    doc.add_paragraph()


def add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bullet-point list."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc: Document, items: list[str]) -> None:
    """Add a numbered list."""
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}) {item}", style="List Number")


def add_bold_paragraph(doc: Document, bold_text: str, normal_text: str = "") -> None:
    """Add a paragraph with a bold prefix."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
